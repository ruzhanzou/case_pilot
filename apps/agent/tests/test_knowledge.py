from pathlib import Path

from docx import Document
from openpyxl import Workbook

from casepilot_agent.knowledge import build_chunks, parse_document, pretokenize
from casepilot_agent.providers.mock import MockProvider
from casepilot_agent.store import lexical_search_terms


def test_markdown_structural_chunking_keeps_parent_child_and_locator(
    tmp_path: Path,
) -> None:
    path = tmp_path / "requirement.md"
    path.write_text(
        "# 账号中心\n## REQ-LOGIN-001 登录\n"
        + "\n".join(f"- 验收规则 {index}：输入合法时返回可观察状态" for index in range(180)),
        encoding="utf-8",
    )
    blocks = parse_document(path, "text/markdown")
    chunks = build_chunks(
        title="账号中心需求",
        blocks=blocks,
        provider=MockProvider(),
    )

    parents = [chunk for chunk in chunks if chunk["chunk_type"] == "parent"]
    children = [chunk for chunk in chunks if chunk["chunk_type"] == "child"]
    assert parents and children
    assert all(child["parent_key"] for child in children)
    assert all(child["token_count"] <= 650 for child in children)
    assert all(len(chunk["embedding"]) == 2048 for chunk in chunks)
    assert "账号中心需求" in children[0]["contextual_content"]


def test_docx_xlsx_and_csv_keep_native_location(tmp_path: Path) -> None:
    docx_path = tmp_path / "rules.docx"
    document = Document()
    document.add_heading("退款规则", level=1)
    document.add_paragraph("重复退款请求必须保持幂等。")
    document.save(docx_path)

    xlsx_path = tmp_path / "cases.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "验收标准"
    sheet.append(["需求编号", "规则"])
    sheet.append(["REQ-REFUND-001", "退款成功后状态可查询"])
    workbook.save(xlsx_path)

    csv_path = tmp_path / "errors.csv"
    csv_path.write_text("错误码,含义\nE1001,订单不存在\n", encoding="utf-8")

    docx_blocks = parse_document(
        docx_path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    xlsx_blocks = parse_document(
        xlsx_path,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    csv_blocks = parse_document(csv_path, "text/csv")

    assert docx_blocks[0].section_path == "退款规则"
    assert "Sheet 验收标准" in xlsx_blocks[0].locator
    assert csv_blocks[0].locator == "行 1-2"


def test_chinese_pretokenization_preserves_requirement_and_error_codes() -> None:
    tokens = pretokenize("REQ-PAY-001 支付失败返回错误码 E1001")
    assert "REQ" in tokens
    assert "E1001" in tokens


def test_chinese_lexical_terms_remove_punctuation_and_keep_question_keywords() -> None:
    terms = lexical_search_terms(
        pretokenize("根据内部知识库，文本摘要默认保留多久，用户能否关闭保存？")
    )

    assert {"文本", "摘要", "默认", "保留", "关闭", "保存"}.issubset(terms)
    assert "，" not in terms
    assert "？" not in terms


def test_chunks_remain_searchable_without_embedding_provider(tmp_path: Path) -> None:
    path = tmp_path / "fallback.md"
    path.write_text("# 支付\nREQ-PAY-001 支付失败返回 E1001", encoding="utf-8")

    chunks = build_chunks(
        title="支付需求",
        blocks=parse_document(path, "text/markdown"),
    )

    assert chunks
    assert all(chunk["embedding"] is None for chunk in chunks)
    assert all(chunk["search_text"] for chunk in chunks)
