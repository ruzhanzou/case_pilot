import csv
import re
from dataclasses import dataclass, field
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

import jieba
import pymupdf as fitz
import pytesseract
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from casepilot_agent.contracts import EmbeddingProvider


@dataclass
class ParsedBlock:
    content: str
    section_path: str = ""
    locator: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


def token_count(text: str) -> int:
    chinese = len(re.findall(r"[\u3400-\u9fff]", text))
    words = len(re.findall(r"[A-Za-z0-9_./:-]+", text))
    punctuation = len(re.findall(r"[^\w\s\u3400-\u9fff]", text))
    return chinese + words + punctuation // 3


def pretokenize(text: str) -> str:
    return " ".join(token.strip() for token in jieba.cut(text) if token.strip())


def _ocr_image(image: Image.Image) -> str:
    return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()


def _parse_pdf(content: bytes) -> list[ParsedBlock]:
    document = fitz.open(stream=content, filetype="pdf")
    blocks: list[ParsedBlock] = []
    for page_index, page in enumerate(document):
        text = page.get_text("text").strip()
        if len(text) < 40:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            text = _ocr_image(image)
        if text:
            blocks.append(
                ParsedBlock(
                    content=text,
                    locator=f"第 {page_index + 1} 页",
                    metadata={"page": page_index + 1},
                )
            )
    return blocks


def _parse_docx(content: bytes) -> list[ParsedBlock]:
    document = Document(BytesIO(content))
    blocks: list[ParsedBlock] = []
    headings: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.lower().startswith("heading"):
            try:
                level = int(paragraph.style.name.split()[-1])
            except ValueError:
                level = 1
            headings = headings[: level - 1] + [text]
            continue
        blocks.append(
            ParsedBlock(
                content=text,
                section_path=" / ".join(headings),
                locator=" / ".join(headings),
            )
        )
    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        if rows:
            blocks.append(
                ParsedBlock(
                    content="\n".join(rows),
                    section_path=" / ".join(headings),
                    locator=f"表格 {table_index}",
                    metadata={"table": table_index},
                )
            )
    return blocks


def _parse_xlsx(content: bytes) -> list[ParsedBlock]:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    blocks: list[ParsedBlock] = []
    for sheet in workbook.worksheets:
        batch: list[str] = []
        start_row = 1
        for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = ["" if value is None else str(value) for value in row]
            if not any(values):
                continue
            if not batch:
                start_row = row_number
            batch.append(" | ".join(values))
            if len(batch) >= 40:
                blocks.append(
                    ParsedBlock(
                        content="\n".join(batch),
                        section_path=sheet.title,
                        locator=f"Sheet {sheet.title}，行 {start_row}-{row_number}",
                        metadata={
                            "sheet": sheet.title,
                            "row_start": start_row,
                            "row_end": row_number,
                        },
                    )
                )
                batch = []
        if batch:
            end_row = start_row + len(batch) - 1
            blocks.append(
                ParsedBlock(
                    content="\n".join(batch),
                    section_path=sheet.title,
                    locator=f"Sheet {sheet.title}，行 {start_row}-{end_row}",
                    metadata={
                        "sheet": sheet.title,
                        "row_start": start_row,
                        "row_end": end_row,
                    },
                )
            )
    return blocks


def _parse_csv(content: bytes) -> list[ParsedBlock]:
    text = content.decode("utf-8-sig", errors="replace")
    rows = list(csv.reader(StringIO(text)))
    blocks: list[ParsedBlock] = []
    for index in range(0, len(rows), 40):
        batch = rows[index : index + 40]
        if not batch:
            continue
        blocks.append(
            ParsedBlock(
                content="\n".join(" | ".join(row) for row in batch),
                locator=f"行 {index + 1}-{index + len(batch)}",
                metadata={"row_start": index + 1, "row_end": index + len(batch)},
            )
        )
    return blocks


def _parse_text(content: bytes) -> list[ParsedBlock]:
    text = content.decode("utf-8-sig", errors="replace")
    headings: list[str] = []
    blocks: list[ParsedBlock] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            blocks.append(
                ParsedBlock(
                    content="\n".join(buffer).strip(),
                    section_path=" / ".join(headings),
                    locator=" / ".join(headings),
                )
            )
            buffer.clear()

    for line in text.splitlines():
        stripped = line.strip()
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        numbered = re.match(r"^(\d+(?:\.\d+)*)[、.\s]+(.+)$", stripped)
        if heading:
            flush()
            level = len(heading.group(1))
            headings = headings[: level - 1] + [heading.group(2)]
        elif numbered and len(stripped) < 120:
            flush()
            level = numbered.group(1).count(".") + 1
            headings = headings[: level - 1] + [stripped]
        elif not stripped:
            flush()
        else:
            buffer.append(stripped)
    flush()
    return [block for block in blocks if block.content]


def parse_document(path: Path, mime_type: str) -> list[ParsedBlock]:
    content = path.read_bytes()
    extension = path.suffix.lower()
    if mime_type == "application/pdf" or extension == ".pdf":
        return _parse_pdf(content)
    if extension == ".docx":
        return _parse_docx(content)
    if extension == ".xlsx":
        return _parse_xlsx(content)
    if extension == ".csv":
        return _parse_csv(content)
    if extension in {".png", ".jpg", ".jpeg"}:
        return [
            ParsedBlock(
                content=_ocr_image(Image.open(BytesIO(content))),
                locator="图片 OCR",
                metadata={"ocr": True},
            )
        ]
    return _parse_text(content)


def _split_text(text: str, max_tokens: int, overlap_tokens: int) -> list[str]:
    if token_count(text) <= max_tokens:
        return [text]
    lines = [line for line in text.splitlines() if line.strip()]
    pieces: list[str] = []
    current: list[str] = []
    for line in lines:
        if current and token_count("\n".join([*current, line])) > max_tokens:
            pieces.append("\n".join(current))
            overlap: list[str] = []
            for previous in reversed(current):
                if token_count("\n".join([previous, *overlap])) > overlap_tokens:
                    break
                overlap.insert(0, previous)
            current = [*overlap, line]
        else:
            current.append(line)
    if current:
        pieces.append("\n".join(current))
    if len(pieces) == 1 and token_count(pieces[0]) > max_tokens:
        char_limit = max(200, max_tokens * 2)
        overlap_chars = overlap_tokens * 2
        return [
            text[index : index + char_limit]
            for index in range(0, len(text), char_limit - overlap_chars)
        ]
    return pieces


def build_chunks(
    *,
    title: str,
    blocks: list[ParsedBlock],
    provider: EmbeddingProvider | None = None,
) -> list[dict[str, Any]]:
    draft_chunks: list[dict[str, Any]] = []
    ordinal = 0
    for block_index, block in enumerate(blocks):
        for parent_index, parent_text in enumerate(
            _split_text(block.content, max_tokens=1800, overlap_tokens=0)
        ):
            parent_key = f"p-{block_index}-{parent_index}"
            context_prefix = "；".join(
                item
                for item in (
                    f"文档：{title}",
                    f"章节：{block.section_path}" if block.section_path else "",
                    f"定位：{block.locator}" if block.locator else "",
                )
                if item
            )
            draft_chunks.append(
                {
                    "key": parent_key,
                    "parent_key": None,
                    "chunk_type": "parent",
                    "ordinal": ordinal,
                    "section_path": block.section_path,
                    "locator": block.locator,
                    "content": parent_text,
                    "contextual_content": f"{context_prefix}\n{parent_text}",
                        "search_text": pretokenize(f"{context_prefix} {parent_text}"),
                        "token_count": token_count(parent_text),
                        "embedding": None,
                        "metadata": block.metadata,
                }
            )
            ordinal += 1
            for child_index, child_text in enumerate(
                _split_text(parent_text, max_tokens=600, overlap_tokens=80)
            ):
                draft_chunks.append(
                    {
                        "key": f"c-{block_index}-{parent_index}-{child_index}",
                        "parent_key": parent_key,
                        "chunk_type": "child",
                        "ordinal": ordinal,
                        "section_path": block.section_path,
                        "locator": block.locator,
                        "content": child_text,
                        "contextual_content": f"{context_prefix}\n{child_text}",
                        "search_text": pretokenize(f"{context_prefix} {child_text}"),
                        "token_count": token_count(child_text),
                        "embedding": None,
                        "metadata": block.metadata,
                    }
                )
                ordinal += 1
    if provider is not None:
        attach_embeddings(draft_chunks, provider)
    return draft_chunks


def attach_embeddings(
    chunks: list[dict[str, Any]],
    provider: EmbeddingProvider,
) -> None:
    vectors = provider.embed(
        [chunk["contextual_content"] for chunk in chunks]
    )
    if len(vectors) != len(chunks):
        raise ValueError("embedding_count_mismatch")
    for chunk, vector in zip(chunks, vectors, strict=True):
        if len(vector) != provider.dimensions:
            raise ValueError("embedding_dimension_mismatch")
        chunk["embedding"] = vector
